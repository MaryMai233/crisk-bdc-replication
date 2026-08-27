from __future__ import annotations

from pathlib import Path
import re
import subprocess
import tempfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
OUTPUT = HERE / "Climate_Risk_and_BDCs_Word.docx"


def set_font(run, size: float, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_border(cell, top=None, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
    for edge, width in (("top", top), ("bottom", bottom)):
        if width:
            element = borders.find(qn(f"w:{edge}"))
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), width)
            element.set(qn("w:color"), "000000")


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r._r.extend([begin, instr, end]); set_font(r, 9)


def replace_paragraph(p, text):
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    r = p.add_run(text); set_font(r, 10.5)


def style_document(doc):
    s = doc.sections[0]
    s.top_margin = Inches(0.80); s.bottom_margin = Inches(0.76)
    s.left_margin = Inches(0.90); s.right_margin = Inches(0.90)
    add_page_number(s)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.06
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    wanted = {"Title": 15, "Heading 1": 12.5, "Heading 2": 11.5}
    for st in doc.styles:
        if st.name not in wanted:
            continue
        size = wanted[st.name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor(0,0,0)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(7); st.paragraph_format.space_after = Pt(3)
    for i,p in enumerate(doc.paragraphs):
        if i < 4: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            if p.style.name not in {"Title","Heading 1","Heading 2"}:
                set_font(r, 10.5, r.bold, r.italic)


def clean_body(doc):
    replacements = {
        "Table [tab:bank]": "Table 1", "Figure [fig:beta]": "Figure 1",
        "Table [tab:bdc]": "Table 2", "Figure [fig:resolution]": "Figure 2",
        "Table [tab:stress]": "Table 3", "[eq:factor]":"(1)", "[eq:crisk]":"(2)", "[eq:coverage]":"(3)",
    }
    for p in doc.paragraphs:
        text=p.text
        for old,new in replacements.items(): text=text.replace(old,new)
        if text.startswith("99 Engle"):
            text = text[3:]
        if text != p.text: replace_paragraph(p,text)


def add_caption(doc, label, title):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True; p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{label}.  {title}"); set_font(r,11,bold=True)


def add_note(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
    r=p.add_run("Notes: "); set_font(r,8,italic=True)
    r=p.add_run(text); set_font(r,8)


def add_table(doc,label,title,rows,header_rows,note,font_size=8.2):
    add_caption(doc,label,title)
    t=doc.add_table(rows=len(rows), cols=max(len(x) for x in rows)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,row in enumerate(rows):
        for j in range(len(t.columns)):
            c=t.cell(i,j); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; c.text=row[j] if j<len(row) else ""
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=1.0
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: set_font(r,font_size,bold=(i in header_rows))
            set_cell_border(c, top="12" if i==0 else None, bottom="8" if i in header_rows or i==len(rows)-1 else None)
    if note: add_note(doc,note)


def add_figure(doc,label,title,image,width=6.2):
    add_caption(doc,label,title)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(str(image), width=Inches(width))


def element_text(el):
    if el.tag != qn("w:p"): return ""
    return "".join(node.text or "" for node in el.iter() if node.tag==qn("w:t")).strip()


def normalize(text): return re.sub(r"\s+"," ",text.replace("\xa0"," ").replace("’", "'")).strip()


def find_anchor(doc,prefix):
    needle=normalize(prefix)
    for p in doc.paragraphs:
        if normalize(p.text).startswith(needle): return p._p
    raise RuntimeError(f"Anchor not found: {prefix}")


def artifact_groups(doc):
    body=doc._element.body; els=[e for e in list(body) if e.tag!=qn("w:sectPr")]
    starts=[]
    for idx,e in enumerate(els):
        tx=element_text(e)
        if re.match(r"^(Table|Figure) \d+\.",tx): starts.append((idx,tx.split(".",1)[0]))
    out={}
    for k,(idx,label) in enumerate(starts):
        end=starts[k+1][0] if k+1<len(starts) else len(els); out[label]=els[idx:end]
    return out


def move_after(anchor, group):
    cur=anchor
    for e in group: cur.addnext(e); cur=e


def append_artifacts(doc):
    add_table(doc,"Table 1","Bank Replication Benchmarks",[
        ["","Replicated","Published","Ratio"],
        ["Mean climate beta, 2019","0.193","—","—"],
        ["Mean climate beta, 2020","0.424","—","—"],
        ["2019–2020 mean change","0.230***","—","—"],
        ["","(7.93)","",""],
        ["Top-four mCRISK, end-2020 (USD bn)","221.7","260.0","0.853"],
        ["Top-four CRISK increase, 2020 (USD bn)","372.8","430.89","0.865"],
    ],{0},"t statistic in parentheses for the paired beta change. ***, **, and * denote p<0.01, p<0.05, and p<0.10.")
    add_figure(doc,"Figure 1","Annual Climate Beta for Banks and BDCs",ROOT/"01_Bank_CRISK_Replication"/"Results"/"Figure_1_Aggregate_Climate_Beta.png",6.15)
    add_table(doc,"Table 2","BDC Climate Beta and Portfolio Climate Beta",[
        ["","(1)","(2)","(3)","(4)"],
        ["Portfolio Climate Beta","0.223*","0.126","0.421***","0.694***"],
        ["","(1.84)","(1.05)","(3.18)","(5.25)"],
        ["N","380","380","380","380"],
        ["BDC Controls","N","Y","Y","Y"],
        ["BDC FE","N","N","Y","Y"],
        ["Year FE","N","N","N","Y"],
        ["Adj. R²","0.012","0.088","0.161","0.356"],
    ],{0},"Quarterly data, 2021Q1–2025Q4. Standard errors are clustered by BDC; t statistics are in parentheses.",8.4)
    add_figure(doc,"Figure 2","Measurement Resolution and the BDC Portfolio Mechanism",ROOT/"02_BDC_Investment_Exposure_and_Climate_Beta"/"Results"/"Figure_3_FF49_DCC_Portfolio_Mechanism.png",5.55)
    add_table(doc,"Table 3","BDC Asset-Coverage Stress",[
        ["Mean reported asset coverage (%)","203.28"],
        ["Mean stressed asset coverage (%)","197.43"],
        ["Mean buffer compression (pp)","5.85"],
        ["Mean compression / mean buffer (%)","11.01"],
        ["Legal breaches","0"],
        ["N","376"],
    ],set(),"",8.5)


def reposition(doc):
    g=artifact_groups(doc)
    a=find_anchor(doc,"The replication recovers the main bank result.")
    move_after(a,g["Figure 1"]); move_after(a,g["Table 1"])
    move_after(find_anchor(doc,"The direct analogue of the original Table 1 is positive"), g["Table 2"])
    move_after(find_anchor(doc,"The remaining precision gap is mainly a measurement issue."), g["Figure 2"])
    move_after(find_anchor(doc,"Under the maintained 50-percent climate-factor stress"), g["Table 3"])


def main():
    with tempfile.TemporaryDirectory(prefix="crisk_word_") as td:
        src=Path(td)/"source.docx"; stripped=Path(td)/"narrative.tex"
        latex=(HERE/"Climate_Risk_and_BDCs.tex").read_text()
        latex=re.sub(r"\\begin\{table\}\[H\].*?\\end\{table\}","",latex,flags=re.S)
        latex=re.sub(r"\\begin\{figure\}\[H\].*?\\end\{figure\}","",latex,flags=re.S)
        stripped.write_text(latex)
        subprocess.run(["pandoc",str(stripped),"--from=latex","--to=docx",f"--resource-path={HERE}:{ROOT}",f"--output={src}"],check=True)
        doc=Document(src); clean_body(doc); style_document(doc); append_artifacts(doc); reposition(doc); doc.save(OUTPUT)
    print(OUTPUT)

if __name__=="__main__": main()
